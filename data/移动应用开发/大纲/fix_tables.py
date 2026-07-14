"""
完善表格内容，突出移动平台工程实践
"""
from docx import Document

doc = Document('软件23+6+《移动应用开发》+教学大纲+20260713.docx')

# ===== 修改表4：平时成绩评价标准 - 课程目标4 =====
table4 = doc.tables[4]
# 行5：课程目标4
table4.rows[5].cells[1].text = '综合开发实践、Git 协作、AI 工具应用'
table4.rows[5].cells[2].text = '独立整合多技术栈完成综合项目，熟练使用 Git 协作与 AI 工具优化，具备移动平台工程实践与性能优化能力。'
table4.rows[5].cells[3].text = '整合多技术栈完成综合项目，能使用 Git 协作与 AI 工具，具备基本的移动平台工程实践能力。'
table4.rows[5].cells[4].text = '基本整合技术栈完成项目，能使用 Git 进行基本版本管理。'
table4.rows[5].cells[5].text = '不能整合技术栈完成项目，不具备协作与优化能力。'
table4.rows[5].cells[6].text = '30'
print('表4修改完成（平时成绩-课程目标4）')

# ===== 修改表5：实验成绩评价标准 - 课程目标4 =====
table5 = doc.tables[5]
# 行5：课程目标4
table5.rows[5].cells[1].text = '跨平台综合项目实战'
table5.rows[5].cells[2].text = '提交可运行的多端应用及完整技术选型对比报告，熟练使用 Git 协作，AI 工具应用深入，具备移动平台工程实践与性能优化能力；实验报告完整规范。'
table5.rows[5].cells[3].text = '提交可运行的多端应用及技术选型报告，能使用 Git 协作与 AI 工具，具备基本的移动平台工程实践能力；实验报告较完整。'
table5.rows[5].cells[4].text = '提交基本可运行的应用及简要报告，能进行基本的 Git 操作。'
table5.rows[5].cells[5].text = '未提交可运行的应用及报告；不能使用 Git 进行协作。'
table5.rows[5].cells[6].text = '30'
print('表5修改完成（实验成绩-课程目标4）')

# ===== 修改表6：期末考核评价内容 =====
table6 = doc.tables[6]
# 行3：课程目标3
table6.rows[3].cells[1].text = '移动平台适配：多端适配方案、权限管理、硬件集成与性能优化'
# 行4：课程目标4
table6.rows[4].cells[1].text = '工程实践能力：AI 工具应用、代码质量、测试验证与应用发布流程'
print('表6修改完成（期末考核）')

# ===== 修改表1：实验项目表 - 实验六学时、实验六名称 =====
table1 = doc.tables[1]
# 行5：实验五（鸿蒙多端应用开发）- 学时改为4
table1.rows[5].cells[3].text = '4'
# 行6：实验六（跨平台项目实战）- 每组人数改为6
table1.rows[6].cells[5].text = '6'
table1.rows[6].cells[2].text = '跨平台综合项目实战'
print('表1修改完成（实验项目表）')

# 保存
doc.save('软件23+6+《移动应用开发》+教学大纲+20260713.docx')
print('\n文件已保存')
