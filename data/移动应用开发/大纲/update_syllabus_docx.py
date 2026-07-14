"""
根据新大纲更新 Word 版本教学大纲
重点修改：第六章综合开发实践、实验六、考核评价标准、教材资源
保持课程目标和毕业要求不变
"""
from docx import Document
from docx.shared import Pt

doc = Document('软件23+6+《移动应用开发》+教学大纲+20260713.docx')

# ===== 修改段落 =====
for i, para in enumerate(doc.paragraphs):
    # 课程简介 - 更新内容
    if i == 14 and '移动应用开发是软件工程的一门专业限选课程' in para.text:
        para.text = '移动应用开发是计算机科学与技术的一门专业限选课程，是一门应用型和实践性很强的课程。本课程的主要内容包括：移动应用开发技术体系（原生/混合/跨平台）、主流开发平台对比、Android/iOS 原生开发基础、跨平台开发框架（Flutter、React Native、Uniapp、MAUI）、微信小程序开发、华为鸿蒙多端应用技术、移动应用与后端交互（RESTful API、JSON 数据解析、Token 认证）、AI 编程工具在移动开发中的应用、UI 设计、数据存储、Git 版本控制与团队协作等。本课程共 48 个学时，其中理论教学为 24 学时、实验教学为 24 学时。实验采用每组 6 人、每人负责一个技术栈的协作探究模式。本课程的考评方式为考查，包括平时考核（20%）、实验考核（30%）和期终考核（50%）。'
        print(f'修改段落{i}: 课程简介')
    
    # 第六章 - 综合开发实践标题
    if '综合开发实践' in para.text and i > 70 and i < 80:
        # 找到第六章的标题位置
        chapter6_start = i
        print(f'找到第六章起始位置: 段落{i}')

# ===== 查找并修改第六章教学内容 =====
# 先找到第六章教学内容的具体位置
for i, para in enumerate(doc.paragraphs):
    if i >= 74 and i <= 83:
        if '项目架构设计：MVP 模式在混合应用中的实践' in para.text:
            # 替换第六章教学内容
            para.text = '项目架构设计：MVP/MVVM 模式在移动应用中的实践、分层架构设计原则。'
            print(f'修改段落{i}: 项目架构设计')
        elif '数据存储方案：SharedPreferences' in para.text:
            para.text = '数据存储方案：SharedPreferences、SQLite、Room（Android）、Core Data（iOS）、小程序本地缓存、鸿蒙数据管理。'
            print(f'修改段落{i}: 数据存储方案')
        elif '性能优化：启动速度优化' in para.text:
            para.text = '移动平台工程实践：权限管理与运行时权限请求、应用生命周期管理、后台任务处理、推送通知集成、蓝牙/WiFi 通信基础。'
            print(f'修改段落{i}: 移动平台工程实践')
        elif 'AI 工具深度应用：利用 AI 工具进行代码重构' in para.text:
            # 在这之前插入几行
            para.text = '性能优化：启动速度优化、内存管理与泄漏检测、网络请求优化、UI 渲染性能调优、图片加载优化。\n移动应用发布流程：应用商店上架规范（App Store、Google Play、华为应用市场）、版本管理策略、用户反馈收集与问题追踪。\n代码质量保障：代码规范、静态分析工具、自动化测试框架（单元测试、集成测试、UI 测试）。\nGit 版本控制与团队协作：分支管理策略、代码审查流程、协作开发规范。\nAI 工具深度应用：利用 AI 工具进行代码重构、性能优化建议分析、测试用例生成。'
            print(f'修改段落{i}: 性能优化+发布流程+代码质量+Git+AI')

# 修改第六章学生学习预期成果
for i, para in enumerate(doc.paragraphs):
    if '整合多技术栈（包括 Xamarin）完成综合项目' in para.text:
        para.text = '学生学习预期成果：整合多技术栈完成综合项目，掌握移动平台工程实践（权限管理、生命周期、性能优化、应用发布），能借助 AI 工具与 Git 进行团队协作开发与优化，具备系统优化能力，支撑课程目标 4。'
        print(f'修改段落{i}: 第六章学习预期成果')

# 修改第六章教学重点
for i, para in enumerate(doc.paragraphs):
    if '教学重点：多技术栈集成方案设计、Xamarin 项目实战。' in para.text:
        para.text = '教学重点：移动平台工程实践、性能优化策略、应用发布流程。'
        print(f'修改段落{i}: 第六章教学重点')

# 修改第六章教学难点
for i, para in enumerate(doc.paragraphs):
    if '教学难点：跨技术栈调试与性能调优。' in para.text and i > 75:
        para.text = '教学难点：跨技术栈调试与性能调优、应用商店上架规范。'
        print(f'修改段落{i}: 第六章教学难点')

# ===== 修改实验部分 =====
# 实验六改为鸿蒙多端应用开发，实验七改为跨平台综合项目实战（增加移动工程实践）
for i, para in enumerate(doc.paragraphs):
    # 修改实验六标题
    if '实验项目六：华为多端应用开发' in para.text:
        para.text = '实验项目六：鸿蒙多端应用开发'
        print(f'修改段落{i}: 实验六标题')
    
    # 修改实验六内容
    if '开发天气应用（手机/平板界面适配）。' in para.text and i > 130:
        para.text = '实验内容：\n使用 DevEco Studio 开发天气应用，实现手机/平板界面自适应布局；通过模拟器演示分布式数据同步原理。扩展案例：调用设备传感器（如光线传感器、陀螺仪）实现简单的物联网数据采集与展示场景。'
        print(f'修改段落{i}: 实验六内容')
    
    # 修改实验六学习预期成果
    if '实现跨设备数据同步，从而支撑课程目标3的达成。' in para.text and i > 130:
        para.text = '学生学习预期成果：\n掌握 ArkUI 声明式语法与多端 UI 适配策略，理解分布式能力原理，了解移动设备传感器与物联网集成基础，从而支撑课程目标 3 的达成。'
        print(f'修改段落{i}: 实验六学习成果')

# 修改实验七（综合项目）- 增加移动工程实践环节
for i, para in enumerate(doc.paragraphs):
    if '团队开发 "校园服务" 应用' in para.text:
        para.text = '实验内容：\n团队开发某业务应用（建议选题涉及传感器或硬件交互场景，如运动健康、智能家居控制等），每人负责一个技术栈（Android/Flutter/React Native/Uniapp/MAUI/微信小程序），使用 Git 进行版本控制与协作，借助 AI 编程工具辅助代码生成与调试。\n移动开发工程实践环节：\n- 权限管理：实现运行时权限请求（如定位、相册、麦克风权限）\n- 生命周期管理：处理应用前后台切换、状态保存与恢复\n- 性能优化：进行启动速度、内存占用、网络请求等性能指标测试与优化\n- 测试验证：编写单元测试用例，进行功能测试与跨端兼容性测试\n- 代码规范：制定团队代码规范，使用静态分析工具检查代码质量\n- 版本管理：建立 Git 分支管理策略，进行代码审查'
        print(f'修改段落{i}: 实验七内容')
    
    if '提交可运行的多端应用及技术选型报告' in para.text and i > 140:
        para.text = '学生学习预期成果：\n提交可运行的多端应用及技术选型对比报告，说明不同框架的应用场景与优势，掌握移动平台工程实践（权限管理、生命周期、性能优化）与 Git 团队协作流程，具备 AI 新范式全流程开发能力，从而支撑课程目标 1-4 的达成。'
        print(f'修改段落{i}: 实验七学习成果')

# ===== 修改教材资源 =====
for i, para in enumerate(doc.paragraphs):
    if '[1] 张引.《Xamarin全栈开发技术与实践》' in para.text:
        para.text = '[1] 张思民. Android Studio 应用程序设计（第2版）[M]. 北京：清华大学出版社，2023.'
        print(f'修改段落{i}: 教材1')
    elif '[2] Ed Burnette著' in para.text:
        para.text = '[2] 黑马程序员. Android 移动开发基础案例教程（第2版）[M]. 北京：人民邮电出版社，2022.'
        print(f'修改段落{i}: 教材2')
    elif '[3] 张思民. Android Studio' in para.text:
        para.text = '[3] 王保明.《Uniapp 跨平台开发实战》[M]. 北京：电子工业出版社，2023.'
        print(f'修改段落{i}: 教材3')
    elif '[4] 黑马程序员.Android移动开发' in para.text:
        para.text = '[4] 亢少军. Flutter 技术入门与实战（第2版）[M]. 北京：机械工业出版社，2022.'
        print(f'修改段落{i}: 教材4')
    elif '[5] 王保明.《Uniapp 跨平台开发实战》' in para.text:
        para.text = '[5] 刘望舒. Android 进阶之光（第2版）[M]. 北京：电子工业出版社，2022.'
        print(f'修改段落{i}: 教材5')
    elif '[6] 亢少军. Flutter 技术入门' in para.text:
        para.text = '[6] 柳峰. 微信小程序开发零基础入门（第2版）[M]. 北京：清华大学出版社，2023.'
        print(f'修改段落{i}: 教材6')
    elif '[7] 刘望舒. Android 进阶之光' in para.text:
        para.text = '[7] 刘长龙. .NET MAUI 跨平台应用开发 [M]. 北京：人民邮电出版社，2023.'
        print(f'修改段落{i}: 教材7')
    elif '[8] 柳峰. 微信小程序' in para.text:
        para.text = '[8] 华为开发者联盟. HarmonyOS 应用开发实战 [M]. 北京：电子工业出版社，2024.'
        print(f'修改段落{i}: 教材8')
    elif '[9] 刘长龙. .NET MAUI' in para.text:
        para.text = '[9] 杜文. Flutter 实战 [M]. 北京：机械工业出版社，2023.'
        print(f'修改段落{i}: 教材9')
    elif '[10] 华为开发者联盟. HarmonyOS' in para.text:
        para.text = '[10] 任玉刚. Android 开发艺术探索 [M]. 北京：电子工业出版社，2022.'
        print(f'修改段落{i}: 教材10')

# 修改制定时间
for i, para in enumerate(doc.paragraphs):
    if '制定时间：2026年2月' in para.text:
        para.text = '制定人（签字）：   审核人（签字）：              制定时间：2026年7月'
        print(f'修改段落{i}: 制定时间')

# ===== 修改实验成绩评定 - 增加实验六和七的描述 =====
# 先看看表格内容
print('\n=== 表格内容 ===')
for ti, table in enumerate(doc.tables):
    print(f'\n表{ti} ({len(table.rows)}行 x {len(table.columns)}列):')
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip()[:30] for cell in row.cells]
        print(f'  行{ri}: {cells}')

# 保存文件
output_file = '软件23+6+《移动应用开发》+教学大纲+20260713.docx'
doc.save(output_file)
print(f'\n文件已保存: {output_file}')
