"""
修复成绩评定表格中的重复内容
"""
from docx import Document

doc = Document('软件23+6+《移动应用开发》+教学大纲+20260713.docx')

# ===== 表4：平时成绩评价标准 =====
table4 = doc.tables[4]

# 行0：表头
table4.rows[0].cells[0].text = '课程目标'
table4.rows[0].cells[1].text = '观测点'
table4.rows[0].cells[2].text = '评价标准'
table4.rows[0].cells[3].text = '评价标准'
table4.rows[0].cells[4].text = '评价标准'
table4.rows[0].cells[5].text = '评价标准'
table4.rows[0].cells[6].text = '成绩比例（%）'

# 行1：子表头
table4.rows[1].cells[0].text = '课程目标'
table4.rows[1].cells[1].text = '观测点'
table4.rows[1].cells[2].text = '优秀\n90-100'
table4.rows[1].cells[3].text = '良好\n70-89'
table4.rows[1].cells[4].text = '合格\n60-69'
table4.rows[1].cells[5].text = '不合格\n0-59'
table4.rows[1].cells[6].text = '成绩比例（%）'

# 行2：课程目标1
table4.rows[2].cells[0].text = '课程目标1（支撑毕业要求 1.4）'
table4.rows[2].cells[1].text = '移动应用开发技术体系、原生开发基础'
table4.rows[2].cells[2].text = '熟练掌握主流工具链及原生开发核心概念，能独立实现页面交互，技术选型分析准确。'
table4.rows[2].cells[3].text = '较好掌握主流工具链，能实现页面交互，技术选型分析基本合理。'
table4.rows[2].cells[4].text = '能配置使用工具，基本实现页面交互，对技术选型有初步认识。'
table4.rows[2].cells[5].text = '不能独立配置工具，无法实现页面交互。'
table4.rows[2].cells[6].text = '15'

# 行3：课程目标2
table4.rows[3].cells[0].text = '课程目标2（支撑毕业要求 3.2）'
table4.rows[3].cells[1].text = '跨平台开发技术、小程序开发、后端 API 交互'
table4.rows[3].cells[2].text = '独立完成跨平台页面开发，熟练运用 RESTful API 实现数据交互，能借助 AI 工具高效开发小程序。'
table4.rows[3].cells[3].text = '完成跨平台页面开发，能运用 RESTful API 实现数据交互，会用 AI 工具辅助小程序开发。'
table4.rows[3].cells[4].text = '基本完成跨平台页面开发，能实现简单的 API 数据交互，了解 AI 工具的基本用法。'
table4.rows[3].cells[5].text = '不能完成跨平台页面开发，无法实现 API 数据交互。'
table4.rows[3].cells[6].text = '25'

# 行4：课程目标3
table4.rows[4].cells[0].text = '课程目标3（支撑毕业要求 4.2）'
table4.rows[4].cells[1].text = '鸿蒙多端应用开发、技术方案对比分析'
table4.rows[4].cells[2].text = '独立使用 DevEco Studio 开发多端应用，能系统对比分析不同技术方案的优劣并形成报告。'
table4.rows[4].cells[3].text = '使用 DevEco Studio 开发多端应用，能对比分析不同技术方案。'
table4.rows[4].cells[4].text = '基本使用 DevEco Studio 开发应用，能进行简单的技术方案对比。'
table4.rows[4].cells[5].text = '不能使用 DevEco Studio 开发应用，无法进行技术方案分析。'
table4.rows[4].cells[6].text = '30'

# 行5：课程目标4
table4.rows[5].cells[0].text = '课程目标4（支撑毕业要求 5.1）'
table4.rows[5].cells[1].text = '综合开发实践、Git 协作、AI 工具应用'
table4.rows[5].cells[2].text = '独立整合多技术栈完成综合项目，熟练使用 Git 协作与 AI 工具优化，具备移动平台工程实践与性能优化能力。'
table4.rows[5].cells[3].text = '整合多技术栈完成综合项目，能使用 Git 协作与 AI 工具，具备基本的移动平台工程实践能力。'
table4.rows[5].cells[4].text = '基本整合技术栈完成项目，能使用 Git 进行基本版本管理。'
table4.rows[5].cells[5].text = '不能整合技术栈完成项目，不具备协作与优化能力。'
table4.rows[5].cells[6].text = '30'

print('表4修复完成（平时成绩评价标准）')

# ===== 表5：实验成绩评价标准 =====
table5 = doc.tables[5]

# 行0：表头
table5.rows[0].cells[0].text = '课程目标'
table5.rows[0].cells[1].text = '观测点'
table5.rows[0].cells[2].text = '评价标准'
table5.rows[0].cells[3].text = '评价标准'
table5.rows[0].cells[4].text = '评价标准'
table5.rows[0].cells[5].text = '评价标准'
table5.rows[0].cells[6].text = '成绩比例（%）'

# 行1：子表头
table5.rows[1].cells[0].text = '课程目标'
table5.rows[1].cells[1].text = '观测点'
table5.rows[1].cells[2].text = '优秀\n90-100'
table5.rows[1].cells[3].text = '良好\n70-89'
table5.rows[1].cells[4].text = '合格\n60-69'
table5.rows[1].cells[5].text = '不合格\n0-59'
table5.rows[1].cells[6].text = '成绩比例（%）'

# 行2：课程目标1
table5.rows[2].cells[0].text = '课程目标1（支撑毕业要求 1.4）'
table5.rows[2].cells[1].text = '开发环境搭建、原生应用开发'
table5.rows[2].cells[2].text = '熟练完成各平台环境配置并运行项目，掌握 Kotlin 开发与 Activity 跳转，组内技术分享内容充实。'
table5.rows[2].cells[3].text = '较好完成环境配置与项目运行，掌握 Activity 跳转与数据传递，能参与组内分享。'
table5.rows[2].cells[4].text = '完成主要平台环境配置，基本掌握 Activity 跳转与数据传递。'
table5.rows[2].cells[5].text = '不能完成环境配置，未掌握 Activity 跳转与数据传递。'
table5.rows[2].cells[6].text = '15'

# 行3：课程目标2
table5.rows[3].cells[0].text = '课程目标2（支撑毕业要求 3.2）'
table5.rows[3].cells[1].text = '跨平台应用开发、微信小程序开发'
table5.rows[3].cells[2].text = '熟练掌握所选框架的异步编程与状态管理，独立实现 RESTful API 数据交互与列表刷新；熟练开发小程序并运用 AI 工具提升效率。'
table5.rows[3].cells[3].text = '较好掌握所选框架异步编程，能实现 API 数据交互与列表刷新；能开发基础功能小程序并使用 AI 工具。'
table5.rows[3].cells[4].text = '基本掌握所选框架用法，能实现简单的数据请求；能开发简单小程序页面。'
table5.rows[3].cells[5].text = '未掌握框架基本用法，不能实现数据请求；不能开发小程序。'
table5.rows[3].cells[6].text = '25'

# 行4：课程目标3
table5.rows[4].cells[0].text = '课程目标3（支撑毕业要求 4.2）'
table5.rows[4].cells[1].text = '鸿蒙多端应用开发'
table5.rows[4].cells[2].text = '熟练使用 DevEco Studio 开发多端应用，实现手机/平板自适应布局，深入理解分布式能力原理并撰写分析报告。'
table5.rows[4].cells[3].text = '较好使用 DevEco Studio 开发应用，实现基本的多端 UI 适配，理解分布式能力原理。'
table5.rows[4].cells[4].text = '基本使用 DevEco Studio 完成应用开发，实现简单的界面布局。'
table5.rows[4].cells[5].text = '不能使用 DevEco Studio 开发应用，未完成界面布局。'
table5.rows[4].cells[6].text = '30'

# 行5：课程目标4
table5.rows[5].cells[0].text = '课程目标4（支撑毕业要求 5.1）'
table5.rows[5].cells[1].text = '跨平台综合项目实战'
table5.rows[5].cells[2].text = '提交可运行的多端应用及完整技术选型对比报告，熟练使用 Git 协作，AI 工具应用深入，具备移动平台工程实践与性能优化能力；实验报告完整规范。'
table5.rows[5].cells[3].text = '提交可运行的多端应用及技术选型报告，能使用 Git 协作与 AI 工具，具备基本的移动平台工程实践能力；实验报告较完整。'
table5.rows[5].cells[4].text = '提交基本可运行的应用及简要报告，能进行基本的 Git 操作。'
table5.rows[5].cells[5].text = '未提交可运行的应用及报告；不能使用 Git 进行协作。'
table5.rows[5].cells[6].text = '30'

print('表5修复完成（实验成绩评价标准）')

# 保存
doc.save('软件23+6+《移动应用开发》+教学大纲+20260713_fixed.docx')
print('\n文件已保存')
