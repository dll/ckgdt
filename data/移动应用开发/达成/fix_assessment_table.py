"""
修改课程过程性考核合理性审核表，使其与软件工程专业新大纲一致
"""
from docx import Document
from docx.shared import Pt

doc = Document('05-移动应用开发+课程过程性考核合理性审核表+刘东良.docx')

# ===== 1. 修改基本信息段落 =====
for para in doc.paragraphs:
    if '2022级计算机科学与技术专业' in para.text:
        para.text = para.text.replace('2022级计算机科学与技术专业', '2023级软件工程专业')
        print(f'修改段落: {para.text}')
    if '表 1  课程支撑毕业要求与课程目标及课程内容的对应关系' in para.text:
        para.text = para.text.replace('表 1', '表1')
        print(f'修改段落: {para.text}')

# ===== 2. 修改表0：课程支撑毕业要求与课程目标及课程内容的对应关系 =====
table0 = doc.tables[0]

# 行1：课程目标1 - 教学内容
table0.rows[1].cells[0].text = '移动应用开发技术体系（原生/混合/跨平台）及主流平台特性，技术选型逻辑，跨平台开发框架和 AI 编程工具的基本使用。'
# 行1：课程目标1 - 课程目标
table0.rows[1].cells[1].text = '掌握移动应用开发技术体系（原生/混合/跨平台）及主流平台特性，理解技术选型逻辑，熟悉跨平台开发框架和 AI 编程工具的基本使用。'
# 行1：课程目标1 - 毕业要求
table0.rows[1].cells[2].text = '1.4 能够将数学、自然科学、工程基础和专业知识综合应用于解决计算机领域复杂工程问题，能够判别计算机系统的复杂性。'

# 行2：课程目标2 - 教学内容
table0.rows[2].cells[0].text = '跨平台开发框架及小程序技术，后端 API 交互，结合 AI 编程工具，设计实现跨平台应用。'
# 行2：课程目标2 - 课程目标
table0.rows[2].cells[1].text = '运用跨平台开发框架及小程序技术，结合 AI 编程工具与后端 API 交互，设计实现跨平台应用，具备需求建模与创新应用能力。'
# 行2：课程目标2 - 毕业要求
table0.rows[2].cells[2].text = '3.2 能够针对特定工程需求设计计算机应用系统，并在设计环节中体现创新意识。'

# 行3：课程目标3 - 教学内容
table0.rows[3].cells[0].text = '多端开发方案对比，鸿蒙多端应用开发，跨设备适配场景分析。'
# 行3：课程目标3 - 课程目标
table0.rows[3].cells[1].text = '调研对比多端开发方案，分析不同技术栈在跨设备适配场景中的优劣，具备技术方案评估与选型能力。'
# 行3：课程目标3 - 毕业要求
table0.rows[3].cells[2].text = '4.2 针对计算机领域复杂工程问题，能够收集、分析与解释已存在的相关产品、模型、系统、方案、开源资料库等资料，并通过信息综合得到合理有效的结论。'

# 行4：课程目标4 - 教学内容
table0.rows[4].cells[0].text = '软件工程规范，现代开发工具（含 AI 编程工具、Git 版本控制），移动平台工程实践，应用测试与优化。'
# 行4：课程目标4 - 课程目标
table0.rows[4].cells[1].text = '遵循软件工程规范，使用现代开发工具（含 AI 编程工具、Git 版本控制）完成应用测试与优化，具备工程实践能力。'
# 行4：课程目标4 - 毕业要求
table0.rows[4].cells[2].text = '5.1 能够运用现代信息技术和工具获取计算机专业重要资料与信息。'

print('表0修改完成')

# ===== 3. 修改表1：课程目标达成情况评价方式及成绩评定对照表 =====
table1 = doc.tables[1]

# 行2：课程目标1
table1.rows[2].cells[0].text = '课程目标1'
table1.rows[2].cells[1].text = '15（20%）'
table1.rows[2].cells[2].text = '15（30%）'
table1.rows[2].cells[3].text = '15（50%）'

# 行3：课程目标2
table1.rows[3].cells[0].text = '课程目标2'
table1.rows[3].cells[1].text = '25（20%）'
table1.rows[3].cells[2].text = '25（30%）'
table1.rows[3].cells[3].text = '25（50%）'

# 行4：课程目标3
table1.rows[4].cells[0].text = '课程目标3'
table1.rows[4].cells[1].text = '30（20%）'
table1.rows[4].cells[2].text = '30（30%）'
table1.rows[4].cells[3].text = '30（50%）'

# 行5：课程目标4
table1.rows[5].cells[0].text = '课程目标4'
table1.rows[5].cells[1].text = '30（20%）'
table1.rows[5].cells[2].text = '30（30%）'
table1.rows[5].cells[3].text = '30（50%）'

print('表1修改完成')

# ===== 4. 修改表2：平时考核 =====
table2 = doc.tables[2]

table2.rows[1].cells[0].text = '课程目标1'
table2.rows[1].cells[1].text = '移动应用开发技术体系及主流平台特性，技术选型逻辑，跨平台开发框架和 AI 编程工具的基本使用。'
table2.rows[1].cells[2].text = '移动应用开发的技术选型能够判别计算机系统的复杂性。'

table2.rows[2].cells[0].text = '课程目标2'
table2.rows[2].cells[1].text = '运用跨平台开发框架及小程序技术，结合 AI 编程工具与后端 API 交互，设计实现跨平台应用，具备需求建模与创新应用能力。'
table2.rows[2].cells[2].text = '跨平台开发框架及小程序技术设计计算机应用系统。'

table2.rows[3].cells[0].text = '课程目标3'
table2.rows[3].cells[1].text = '调研对比多端开发方案，分析不同技术栈在跨设备适配场景中的优劣，具备技术方案评估与选型能力。'
table2.rows[3].cells[2].text = '多端开发方案对比分析通过信息综合得到合理有效的结论。'

table2.rows[4].cells[0].text = '课程目标4'
table2.rows[4].cells[1].text = '遵循软件工程规范，使用现代开发工具完成应用测试与优化。'
table2.rows[4].cells[2].text = '移动应用工程化开发应用，具备工程实践能力。'

print('表2修改完成')

# ===== 5. 修改表3：实验考核 =====
table3 = doc.tables[3]

table3.rows[1].cells[0].text = '课程目标1'
table3.rows[1].cells[1].text = '移动应用开发技术体系及主流平台特性，跨平台开发框架和 AI 编程工具的基本使用。'
table3.rows[1].cells[2].text = '移动应用开发的技术选型能够判别计算机系统的复杂性。'

table3.rows[2].cells[0].text = '课程目标2'
table3.rows[2].cells[1].text = '运用跨平台开发框架及小程序技术，结合 AI 编程工具与后端 API 交互，设计实现跨平台应用，具备需求建模与创新应用能力。'
table3.rows[2].cells[2].text = '跨平台开发框架及小程序技术设计计算机应用系统。'

table3.rows[3].cells[0].text = '课程目标3'
table3.rows[3].cells[1].text = '调研对比多端开发方案，鸿蒙多端应用开发，跨设备适配，具备技术方案评估与选型能力。'
table3.rows[3].cells[2].text = '多端应用开发与方案对比通过信息综合得到合理有效的结论。'

table3.rows[4].cells[0].text = '课程目标4'
table3.rows[4].cells[1].text = '移动平台工程实践（权限管理、生命周期、性能优化），Git 协作，AI 工具应用，应用测试与优化。'
table3.rows[4].cells[2].text = '移动应用工程化开发与性能优化，具备工程实践能力。'

print('表3修改完成')

# ===== 6. 修改表4：期末考试考核 =====
table4 = doc.tables[4]

table4.rows[1].cells[0].text = '课程目标1'
table4.rows[1].cells[1].text = '移动应用开发技术体系及主流平台特性，技术选型逻辑。'
table4.rows[1].cells[2].text = '移动应用开发的技术选型能够判别计算机系统的复杂性。'

table4.rows[2].cells[0].text = '课程目标2'
table4.rows[2].cells[1].text = '跨平台开发框架及小程序技术，结合 AI 编程工具，设计实现跨平台应用。'
table4.rows[2].cells[2].text = '跨平台开发框架及小程序技术设计计算机应用系统，并创新应用。'

table4.rows[3].cells[0].text = '课程目标3'
table4.rows[3].cells[1].text = '多端开发方案对比，鸿蒙多端应用开发，跨设备适配与技术选型。'
table4.rows[3].cells[2].text = '多端开发方案对比分析通过信息综合得到合理有效的结论。'

table4.rows[4].cells[0].text = '课程目标4'
table4.rows[4].cells[1].text = '移动平台工程实践，性能优化，软件工程规范，应用测试与发布流程。'
table4.rows[4].cells[2].text = '移动应用工程化开发与质量保障，具备工程实践能力。'

print('表4修改完成')

# 保存文件
output_file = '05-移动应用开发+课程过程性考核合理性审核表+刘东良-软件.docx'
doc.save(output_file)
print(f'\n文件已保存: {output_file}')
