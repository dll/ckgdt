"""
更新课程目标达成评价报告
"""
from docx import Document
from docx.shared import Pt

doc = Document('10-课程目标达成评价报告-样例.docx')

# ===== 修改标题段落 =====
doc.paragraphs[0].text = '2026-2027学年第2学期《软件工程》课程目标达成'
doc.paragraphs[1].text = '评价报告'

# ===== 表0：基本信息 =====
table0 = doc.tables[0]
# 行1：课程名称等
table0.rows[1].cells[1].text = '软件工程'
table0.rows[1].cells[4].text = '专业限选课'
table0.rows[1].cells[6].text = '48(24+24)'
table0.rows[1].cells[8].text = '2.5'

# 行2：班级等
table0.rows[2].cells[1].text = 'CKGDT 2026 春季班'
table0.rows[2].cells[4].text = '39'
table0.rows[2].cells[6].text = '206004'
table0.rows[2].cells[8].text = '讲师'

# 行3：命题教师等
table0.rows[3].cells[1].text = '206004'
table0.rows[3].cells[4].text = '考查'
table0.rows[3].cells[6].text = '大作业'

# 行4：阅卷教师等
table0.rows[4].cells[1].text = '206004'
table0.rows[4].cells[5].text = '个人阅卷'

# 行5：评价方法
table0.rows[5].cells[1].text = '定量评价+定性评价'

# 行7-11：课程支撑毕业要求与课程目标对应关系
# 行9：表头
table0.rows[8].cells[0].text = '支撑的毕业要求'
table0.rows[8].cells[1].text = '支撑的毕业要求'
table0.rows[8].cells[2].text = '支撑的毕业要求'
table0.rows[8].cells[3].text = '课程目标'
table0.rows[8].cells[4].text = '课程目标'
table0.rows[8].cells[5].text = '考核内容'
table0.rows[8].cells[6].text = '考核内容'
table0.rows[8].cells[7].text = '考核内容'
table0.rows[8].cells[8].text = '考核内容'

# 行9：课程目标1
table0.rows[9].cells[0].text = '指标点2.2'
table0.rows[9].cells[1].text = '指标点2.2'
table0.rows[9].cells[2].text = '指标点2.2'
table0.rows[9].cells[3].text = '课程目标1'
table0.rows[9].cells[4].text = '课程目标1'
table0.rows[9].cells[5].text = '空间软件维护与项目管理;空间信息系统安全与伦理。'
table0.rows[9].cells[6].text = '空间软件维护与项目管理;空间信息系统安全与伦理。'
table0.rows[9].cells[7].text = '空间软件维护与项目管理;空间信息系统安全与伦理。'
table0.rows[9].cells[8].text = '空间软件维护与项目管理;空间信息系统安全与伦理。'

# 行10：课程目标2
table0.rows[10].cells[0].text = '指标点3.1'
table0.rows[10].cells[1].text = '指标点3.1'
table0.rows[10].cells[2].text = '指标点3.1'
table0.rows[10].cells[3].text = '课程目标2'
table0.rows[10].cells[4].text = '课程目标2'
table0.rows[10].cells[5].text = '空间结构化分析与设计;空间面向对象与EA UML建模;空间面向对象的分析与设计;空间编码与AI辅助开发;空间软件测试与AI辅助验证。'
table0.rows[10].cells[6].text = '空间结构化分析与设计;空间面向对象与EA UML建模;空间面向对象的分析与设计;空间编码与AI辅助开发;空间软件测试与AI辅助验证。'
table0.rows[10].cells[7].text = '空间结构化分析与设计;空间面向对象与EA UML建模;空间面向对象的分析与设计;空间编码与AI辅助开发;空间软件测试与AI辅助验证。'
table0.rows[10].cells[8].text = '空间结构化分析与设计;空间面向对象与EA UML建模;空间面向对象的分析与设计;空间编码与AI辅助开发;空间软件测试与AI辅助验证。'

# 行11：课程目标3
table0.rows[11].cells[0].text = '指标点11.1'
table0.rows[11].cells[1].text = '指标点11.1'
table0.rows[11].cells[2].text = '指标点11.1'
table0.rows[11].cells[3].text = '课程目标3'
table0.rows[11].cells[4].text = '课程目标3'
table0.rows[11].cells[5].text = '空间需求工程与AI辅助分析。'
table0.rows[11].cells[6].text = '空间需求工程与AI辅助分析。'
table0.rows[11].cells[7].text = '空间需求工程与AI辅助分析。'
table0.rows[11].cells[8].text = '空间需求工程与AI辅助分析。'

print('表0修改完成')

# ===== 表1：课程目标达成考核与评价方式及成绩评定对照表 =====
table1 = doc.tables[1]

# 行2：课程目标1
table1.rows[2].cells[0].text = '课程目标1'
table1.rows[2].cells[1].text = '0.20'
table1.rows[2].cells[2].text = '指标点2.2'
table1.rows[2].cells[3].text = '20（20%）'
table1.rows[2].cells[4].text = '20（30%）'
table1.rows[2].cells[5].text = '20（50%）'

# 行3：课程目标2
table1.rows[3].cells[0].text = '课程目标2'
table1.rows[3].cells[1].text = '0.50'
table1.rows[3].cells[2].text = '指标点3.1'
table1.rows[3].cells[3].text = '50（20%）'
table1.rows[3].cells[4].text = '50（30%）'
table1.rows[3].cells[5].text = '50（50%）'

# 行4：课程目标3
table1.rows[4].cells[0].text = '课程目标3'
table1.rows[4].cells[1].text = '0.30'
table1.rows[4].cells[2].text = '指标点11.1'
table1.rows[4].cells[3].text = '30（20%）'
table1.rows[4].cells[4].text = '30（30%）'
table1.rows[4].cells[5].text = '30（50%）'

print('表1修改完成')

# ===== 表2：平时成绩评价标准 =====
table2 = doc.tables[2]

# 行2：课程目标1
table2.rows[2].cells[0].text = '课程目标1'
table2.rows[2].cells[1].text = '空间软件维护与项目管理;空间信息系统安全与伦理。'
table2.rows[2].cells[2].text = '全面掌握，表现突出'
table2.rows[2].cells[3].text = '较好掌握，表现良好'
table2.rows[2].cells[4].text = '基本掌握，表现一般'
table2.rows[2].cells[5].text = '未能掌握，需要改进'
table2.rows[2].cells[6].text = '20'

# 行3：课程目标2
table2.rows[3].cells[0].text = '课程目标2'
table2.rows[3].cells[1].text = '空间结构化分析与设计;空间面向对象与EA UML建模;空间面向对象的分析与设计;空间编码与AI辅助开发;空间软件测试与AI辅助验证。'
table2.rows[3].cells[2].text = '全面掌握，表现突出'
table2.rows[3].cells[3].text = '较好掌握，表现良好'
table2.rows[3].cells[4].text = '基本掌握，表现一般'
table2.rows[3].cells[5].text = '未能掌握，需要改进'
table2.rows[3].cells[6].text = '50'

# 行4：课程目标3
table2.rows[4].cells[0].text = '课程目标3'
table2.rows[4].cells[1].text = '空间需求工程与AI辅助分析。'
table2.rows[4].cells[2].text = '全面掌握，表现突出'
table2.rows[4].cells[3].text = '较好掌握，表现良好'
table2.rows[4].cells[4].text = '基本掌握，表现一般'
table2.rows[4].cells[5].text = '未能掌握，需要改进'
table2.rows[4].cells[6].text = '30'

print('表2修改完成')

# ===== 表3：实验成绩评价标准 =====
table3 = doc.tables[3]

# 行2：课程目标1
table3.rows[2].cells[0].text = '课程目标1'
table3.rows[2].cells[1].text = '空间软件维护与项目管理;空间信息系统安全与伦理。'
table3.rows[2].cells[2].text = '独立完成，结果正确'
table3.rows[2].cells[3].text = '基本完成，结果较好'
table3.rows[2].cells[4].text = '能够完成，有少量错误'
table3.rows[2].cells[5].text = '无法完成或错误较多'
table3.rows[2].cells[6].text = '20'

# 行3：课程目标2
table3.rows[3].cells[0].text = '课程目标2'
table3.rows[3].cells[1].text = '空间结构化分析与设计;空间面向对象与EA UML建模;空间面向对象的分析与设计;空间编码与AI辅助开发;空间软件测试与AI辅助验证。'
table3.rows[3].cells[2].text = '独立完成，结果正确'
table3.rows[3].cells[3].text = '基本完成，结果较好'
table3.rows[3].cells[4].text = '能够完成，有少量错误'
table3.rows[3].cells[5].text = '无法完成或错误较多'
table3.rows[3].cells[6].text = '50'

# 行4：课程目标3
table3.rows[4].cells[0].text = '课程目标3'
table3.rows[4].cells[1].text = '空间需求工程与AI辅助分析。'
table3.rows[4].cells[2].text = '独立完成，结果正确'
table3.rows[4].cells[3].text = '基本完成，结果较好'
table3.rows[4].cells[4].text = '能够完成，有少量错误'
table3.rows[4].cells[5].text = '无法完成或错误较多'
table3.rows[4].cells[6].text = '30'

print('表3修改完成')

# ===== 表4：期末考核评价内容 =====
table4 = doc.tables[4]

# 行1：课程目标1
table4.rows[1].cells[0].text = '课程目标1'
table4.rows[1].cells[1].text = '空间软件维护与项目管理;空间信息系统安全与伦理。'
table4.rows[1].cells[2].text = '20'

# 行2：课程目标2
table4.rows[2].cells[0].text = '课程目标2'
table4.rows[2].cells[1].text = '空间结构化分析与设计;空间面向对象与EA UML建模;空间面向对象的分析与设计;空间编码与AI辅助开发;空间软件测试与AI辅助验证。'
table4.rows[2].cells[2].text = '50'

# 行3：课程目标3
table4.rows[3].cells[0].text = '课程目标3'
table4.rows[3].cells[1].text = '空间需求工程与AI辅助分析。'
table4.rows[3].cells[2].text = '30'

print('表4修改完成')

# ===== 表5：达成度计算表 =====
table5 = doc.tables[5]

# 行5-7：目标1 平时/实验/考核
table5.rows[5].cells[0].text = '目标1'
table5.rows[5].cells[1].text = '0.20'
table5.rows[5].cells[2].text = '平时成绩'
table5.rows[5].cells[3].text = '20'
table5.rows[5].cells[4].text = '16.74'
table5.rows[5].cells[5].text = '0.8368'
table5.rows[5].cells[6].text = '0.20'
table5.rows[5].cells[7].text = '0.8408'
table5.rows[5].cells[8].text = '指标点2.2'
table5.rows[5].cells[9].text = '0.8408'

table5.rows[6].cells[0].text = '目标1'
table5.rows[6].cells[2].text = '实验成绩'
table5.rows[6].cells[3].text = '20'
table5.rows[6].cells[4].text = '14.45'
table5.rows[6].cells[5].text = '0.7226'
table5.rows[6].cells[6].text = '0.30'

table5.rows[7].cells[0].text = '目标1'
table5.rows[7].cells[2].text = '考核成绩'
table5.rows[7].cells[3].text = '20'
table5.rows[7].cells[4].text = '18.27'
table5.rows[7].cells[5].text = '0.9133'
table5.rows[7].cells[6].text = '0.50'

# 行8-10：目标2
table5.rows[8].cells[0].text = '目标2'
table5.rows[8].cells[1].text = '0.50'
table5.rows[8].cells[2].text = '平时成绩'
table5.rows[8].cells[3].text = '50'
table5.rows[8].cells[4].text = '35.37'
table5.rows[8].cells[5].text = '0.7074'
table5.rows[8].cells[6].text = '0.20'
table5.rows[8].cells[7].text = '0.8269'
table5.rows[8].cells[8].text = '指标点3.1'
table5.rows[8].cells[9].text = '0.8269'

table5.rows[9].cells[0].text = '目标2'
table5.rows[9].cells[2].text = '实验成绩'
table5.rows[9].cells[3].text = '50'
table5.rows[9].cells[4].text = '40.10'
table5.rows[9].cells[5].text = '0.8019'
table5.rows[9].cells[6].text = '0.30'

table5.rows[10].cells[0].text = '目标2'
table5.rows[10].cells[2].text = '考核成绩'
table5.rows[10].cells[3].text = '50'
table5.rows[10].cells[4].text = '44.49'
table5.rows[10].cells[5].text = '0.8897'
table5.rows[10].cells[6].text = '0.50'

# 行11-13：目标3
table5.rows[11].cells[0].text = '目标3'
table5.rows[11].cells[1].text = '0.30'
table5.rows[11].cells[2].text = '平时成绩'
table5.rows[11].cells[3].text = '30'
table5.rows[11].cells[4].text = '0.00'
table5.rows[11].cells[5].text = '0.0000'
table5.rows[11].cells[6].text = '0.20'
table5.rows[11].cells[7].text = '0.6577'
table5.rows[11].cells[8].text = '指标点11.1'
table5.rows[11].cells[9].text = '0.6577'

table5.rows[12].cells[0].text = '目标3'
table5.rows[12].cells[2].text = '实验成绩'
table5.rows[12].cells[3].text = '30'
table5.rows[12].cells[4].text = '22.15'
table5.rows[12].cells[5].text = '0.7385'
table5.rows[12].cells[6].text = '0.30'

table5.rows[13].cells[0].text = '目标3'
table5.rows[13].cells[2].text = '考核成绩'
table5.rows[13].cells[3].text = '30'
table5.rows[13].cells[4].text = '26.17'
table5.rows[13].cells[5].text = '0.8723'
table5.rows[13].cells[6].text = '0.50'

# 行14：总体达成度
table5.rows[14].cells[1].text = '0.6'
table5.rows[14].cells[2].text = '0.6'
table5.rows[14].cells[3].text = '0.6'
table5.rows[14].cells[4].text = '0.6'
table5.rows[14].cells[5].text = '课程总体目标达成度'
table5.rows[14].cells[6].text = '0.7789'

print('表5修改完成')

# ===== 表6：达成结果分析 =====
table6 = doc.tables[6]

# 行1：定量评价 - 更新内容
table6.rows[1].cells[1].text = '课程目标1达成度：0.8408（达成）\n课程目标2达成度：0.8269（达成）\n课程目标3达成度：0.6577（达成）\n课程总体达成度：0.7789（达成）\n\n图1 学生个体课程目标1达成情况评价结果\n图2 学生个体课程目标2达成情况评价结果\n图3 学生个体课程目标3达成情况评价结果'

# 行2：定性评价
table6.rows[2].cells[1].text = '1. 调查问卷内容\n调查说明\n本次问卷旨在评估本课程目标对毕业要求的支撑效果，请同学们根据实际情况如实填写。\n\n共回收有效问卷 40 份，综合满意度为 0.0%。\n\n2. 调查结果分析\n（根据实际问卷数据填写）'

# 行3：达成情况分析及持续改进
table6.rows[3].cells[1].text = '''一、定量评价情况分析
本课程设置 3 项课程目标，分别对应毕业要求指标点 2.2、3.1、11.1。

课程目标1（达成度：0.8408，良好，大部分学生达到预期）
课程目标1主要考核掌握通用软件工程基础理论、软件生命周期与主流过程模型。该目标通过平时成绩（20%）、实验成绩（30%）、考核成绩（50%）综合评定。
- 平时环节达成度：0.8368
- 实验环节达成度：0.7226
- 考核环节达成度：0.9133

课程目标2（达成度：0.8269，良好，大部分学生达到预期）
课程目标2主要考核系统掌握空间信息系统需求分析、结构化与面向对象建模、架构设计。该目标通过平时成绩（20%）、实验成绩（30%）、考核成绩（50%）综合评定。
- 平时环节达成度：0.7074
- 实验环节达成度：0.8019
- 考核环节达成度：0.8897
- 有 1 名学生该目标达成度低于0.60，需个别辅导

课程目标3（达成度：0.6577，达标但有提升空间）
课程目标3主要考核理解空间信息类软件工程项目管理、成本管控、进度规划与风险预警。该目标通过平时成绩（20%）、实验成绩（30%）、考核成绩（50%）综合评定。
- 平时环节达成度：0.0000
- 实验环节达成度：0.7385
- 考核环节达成度：0.8723
- 有 3 名学生该目标达成度低于0.60，需个别辅导

二、定性评价情况分析
共回收有效问卷 40 份，综合满意度为 0.0%。

三、教学持续改进

本轮教学改进措施执行情况：
1. 在平时作业中加大与课程目标相关的分析应用问题的题目训练
2. 在每一章结束后，在作业中增加与该章知识点相关的文献阅读培训
3. 根据大纲对照表复核各考核环节比例，优化过程性评价和终结性评价的衔接

后续教学持续改进措施：
1. 课程目标1（0.8408）：保持现有教学节奏，适当提高考核难度
2. 课程目标2（0.8269）：保持现有教学节奏，适当提高考核难度
3. 课程目标3（0.6577）：加大「理解空间信息类软件工程项目管理」相关的对比分析训练，补充测验题目'''

# 行4：签字日期
table6.rows[4].cells[4].text = '2026-07-14'

print('表6修改完成')

# 保存文件
output_file = '10-课程目标达成评价报告-CKGDT2026.docx'
doc.save(output_file)
print(f'\n文件已保存: {output_file}')
